"""
parse_questionnaire.py
----------------------
Parses a questionnaire design file (docx, xlsx, pdf, txt/md) and writes
a structured questionnaire-design.md following the format defined in SKILL.md.

Usage:
    python parse_questionnaire.py <input_file> [output_file]

If output_file is omitted, writes questionnaire-design.md next to the input.

Dependencies (install what you need):
    pip install python-docx openpyxl pdfplumber polars
"""

import sys
import re
from pathlib import Path


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _clean(text: str) -> str:
    return text.strip() if text else ""


def _write_md(questions: list[dict], output_path: Path) -> None:
    lines = ["# QUESTIONNAIRE DESIGN\n"]
    for q in questions:
        lines.append("---")
        lines.append(f"Question: {q['id']}")
        if q.get("type"):
            lines.append(f"Type: {q['type']}")
        if q.get("label"):
            lines.append(f"Label / Text: {q['label']}")
        options = q.get("options", [])
        if options:
            lines.append("Options:")
            for opt in options:
                lines.append(f"  {opt}")
        if q.get("logic"):
            lines.append(f"Logic: {q['logic']}")
        if q.get("note"):
            lines.append(f"Note: {q['note']}")
        lines.append("")  # blank line between blocks

    output_path.write_text("\n".join(lines), encoding="utf-8")
    print(f"Written: {output_path}")


def _summarise(questions: list[dict]) -> None:
    print(f"\nTotal questions: {len(questions)}")
    flagged = [q["id"] for q in questions if q.get("note")]
    if flagged:
        print(f"Questions with ambiguous logic (Note:): {', '.join(flagged)}")
    else:
        print("All routing logic resolved cleanly.")


# ---------------------------------------------------------------------------
# Format-specific parsers
# ---------------------------------------------------------------------------


def _parse_txt(path: Path) -> list[dict]:
    """
    Minimal plain-text / markdown parser.

    Expects blocks separated by blank lines or '---', starting with a line
    like 'Q1', 'Q1.', '1.', or 'Question: Q1'.
    """
    text = path.read_text(encoding="utf-8")
    # Split on blank lines or explicit separators
    raw_blocks = re.split(r"\n(?:---+|\s*)\n", text)

    questions = []
    q_counter = 0

    for block in raw_blocks:
        block = block.strip()
        if not block:
            continue

        lines = [line.rstrip() for line in block.splitlines()]

        # Try to detect question ID on the first non-empty line
        first = lines[0] if lines else ""
        id_match = re.match(r"^(?:Question:\s*)?([A-Za-z]?\d+[a-z]?)[.:\s]", first)
        if not id_match:
            continue  # not a question block

        q_counter += 1
        qid = id_match.group(1).upper()
        label = re.sub(
            r"^(?:Question:\s*)?[A-Za-z]?\d+[a-z]?[.:\s]*", "", first
        ).strip()

        options = []
        logic_parts = []
        note = ""

        for line in lines[1:]:
            # Answer options: lines starting with a number
            opt_match = re.match(r"^(\d+)[.)]\s+(.+)", line)
            if opt_match:
                options.append(f"{opt_match.group(1)}. {opt_match.group(2)}")
                continue
            # Logic hints
            if re.search(
                r"\b(ask|skip|show|if|all respondents|logic|routing)\b", line, re.I
            ):
                logic_parts.append(line.strip())
                continue
            # Extra label text
            if label and line:
                label += " " + line.strip()
            elif line:
                label = line.strip()

        logic = "; ".join(logic_parts) if logic_parts else "All respondents"

        questions.append(
            {
                "id": qid,
                "type": _infer_type(options),
                "label": label,
                "options": options,
                "logic": logic,
                "note": note,
            }
        )

    return questions


def _parse_docx(path: Path) -> list[dict]:
    try:
        from docx import Document  # type: ignore
    except ImportError:
        sys.exit("python-docx not installed. Run: pip install python-docx")

    doc = Document(str(path))
    questions = []

    # --- Try table-based layout first ---
    if doc.tables:
        for table in doc.tables:
            headers = [_clean(c.text).lower() for c in table.rows[0].cells]
            # Detect column positions
            id_col = _find_col(headers, ["q", "id", "code", "no", "number"])
            label_col = _find_col(headers, ["label", "text", "question", "wording"])
            type_col = _find_col(headers, ["type", "format"])
            options_col = _find_col(headers, ["option", "answer", "choice", "response"])
            logic_col = _find_col(
                headers, ["logic", "routing", "skip", "condition", "filter"]
            )

            for row in table.rows[1:]:
                cells = [_clean(c.text) for c in row.cells]
                if not any(cells):
                    continue
                qid = cells[id_col] if id_col is not None else f"Q{len(questions) + 1}"
                label = cells[label_col] if label_col is not None else ""
                qtype = cells[type_col] if type_col is not None else ""
                raw_opts = cells[options_col] if options_col is not None else ""
                logic = cells[logic_col] if logic_col is not None else "All respondents"

                options = _parse_options_text(raw_opts)
                questions.append(
                    {
                        "id": qid.upper(),
                        "type": qtype or _infer_type(options),
                        "label": label,
                        "options": options,
                        "logic": logic or "All respondents",
                        "note": "",
                    }
                )
        if questions:
            return questions

    # --- Fallback: paragraph-based layout ---
    return _parse_txt(Path(str(path)))  # reuse txt parser on extracted text


def _parse_xlsx(path: Path) -> list[dict]:
    try:
        import openpyxl  # type: ignore
    except ImportError:
        sys.exit("openpyxl not installed. Run: pip install openpyxl")

    wb = openpyxl.load_workbook(str(path), data_only=True)
    ws = wb.active

    if ws is None:
        raise FileNotFoundError("None of active worksheet found")

    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        return []

    # Detect header row (first row with >2 non-empty cells)
    header_idx = 0
    for i, row in enumerate(rows):
        non_empty = [str(c).lower() for c in row if c is not None]
        if len(non_empty) >= 2:
            header_idx = i
            break

    headers = [str(c).lower() if c is not None else "" for c in rows[header_idx]]
    id_col = _find_col(headers, ["q", "id", "code", "no", "number"])
    label_col = _find_col(headers, ["label", "text", "question", "wording"])
    type_col = _find_col(headers, ["type", "format"])
    options_col = _find_col(headers, ["option", "answer", "choice", "response"])
    logic_col = _find_col(headers, ["logic", "routing", "skip", "condition", "filter"])

    questions = []
    for row in rows[header_idx + 1 :]:
        cells = [str(c).strip() if c is not None else "" for c in row]
        if not any(cells):
            continue
        qid = cells[id_col] if id_col is not None else f"Q{len(questions) + 1}"
        label = cells[label_col] if label_col is not None else ""
        qtype = cells[type_col] if type_col is not None else ""
        raw_opts = cells[options_col] if options_col is not None else ""
        logic = cells[logic_col] if logic_col is not None else "All respondents"

        if not qid:
            continue

        options = _parse_options_text(raw_opts)
        questions.append(
            {
                "id": qid.upper(),
                "type": qtype or _infer_type(options),
                "label": label,
                "options": options,
                "logic": logic or "All respondents",
                "note": "",
            }
        )

    return questions


def _parse_pdf(path: Path) -> list[dict]:
    try:
        import pdfplumber  # type: ignore
    except ImportError:
        sys.exit("pdfplumber not installed. Run: pip install pdfplumber")

    full_text_lines = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            # Try table extraction first
            tables = page.extract_tables()
            if tables:
                for table in tables:
                    for row in table:
                        full_text_lines.append(
                            "\t".join(str(c) if c else "" for c in row)
                        )
            else:
                text = page.extract_text()
                if text:
                    full_text_lines.extend(text.splitlines())

    # Write to a temp txt and reuse txt parser
    tmp = path.with_suffix(".pdf_extracted.txt")
    tmp.write_text("\n".join(full_text_lines), encoding="utf-8")
    result = _parse_txt(tmp)
    tmp.unlink(missing_ok=True)
    return result


# ---------------------------------------------------------------------------
# Utility helpers
# ---------------------------------------------------------------------------


def _find_col(headers: list[str], keywords: list[str]) -> int | None:
    for kw in keywords:
        for i, h in enumerate(headers):
            if kw in h:
                return i
    return None


def _parse_options_text(raw: str) -> list[str]:
    """Parse an options cell or string into a list of '1. Label' strings."""
    if not raw:
        return []
    options = []
    # Lines like "1. Male\n2. Female" or "1-Male, 2-Female"
    for line in re.split(r"[\n;|]", raw):
        line = line.strip()
        m = re.match(r"^(\d+)[.)\-\s]+(.+)", line)
        if m:
            options.append(f"{m.group(1)}. {m.group(2).strip()}")
        elif line:
            options.append(line)
    return options


def _infer_type(options: list[str]) -> str:
    if not options:
        return "Open"
    return "Single"


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------


def parse(input_path: str, output_path: str | None = None) -> None:
    src = Path(input_path)
    if not src.exists():
        sys.exit(f"File not found: {src}")

    ext = src.suffix.lower()
    if ext == ".docx":
        questions = _parse_docx(src)
    elif ext in (".xlsx", ".xls"):
        questions = _parse_xlsx(src)
    elif ext == ".pdf":
        questions = _parse_pdf(src)
    elif ext in (".txt", ".md", ".csv"):
        questions = _parse_txt(src)
    else:
        sys.exit(f"Unsupported format: {ext}. Supported: .docx .xlsx .pdf .txt .md")

    if not questions:
        print("No questions detected. Check the file layout and adjust the parser.")
        return

    dest = Path(output_path) if output_path else src.parent / "questionnaire-design.md"
    _write_md(questions, dest)
    _summarise(questions)


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)
    parse(sys.argv[1], sys.argv[2] if len(sys.argv) > 2 else None)
